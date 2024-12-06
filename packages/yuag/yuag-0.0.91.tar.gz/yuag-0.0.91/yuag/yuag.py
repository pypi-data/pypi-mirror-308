# استدعاء المكتبات
def PyPDF2():
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    import PyPDF2
    return PyPDF2

def Translator():
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    from googletrans import Translator
    from langdetect import detect

    return Translator

def openai():
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    import openai

    return openai

def BeautifulSoup():
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    from bs4 import BeautifulSoup
    return BeautifulSoup

def requests():
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    import requests
    return requests

def time():
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    import time
    return time

def os():
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    import os
    return os

# global defs
def wait(wait_time: int = 1):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    import time

    time.sleep(wait_time)

def clear():
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    مسح الطَّرَفِيَّة
    """

    import os

    os.system("cls" if os.name == "nt" else "clear")

def detect_lang(text: str, detectOrSrc: int = 0):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    from googletrans import Translator
    from langdetect import detect
    
    try:
        if detectOrSrc == 0: return detect(text)
        else: return Translator().translate(text).src
    except:
        return 404

def rangeNum(num):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------

    ```
    rangeNum(3)

    ==> [0, 1, 2]
    ```
    """
    
    return list(range(num))

def python_path():
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    ارجاع قيمة مسار بايثون، مثل:
    `C:\\Python\\python.exe`
    """
    
    import sys

    return sys.executable

def check_lib_status(lib_to_check: str|list, notFound_return = "not found", installed_lib_return = "installed", builtIn_return = "built-in"):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    التأكد من وجود أي مكتبة أو أكثر من مكتبة في بايثون\n\n

    مكتبة واحدة:
    ------------
    ```
    lib_to_check = "إسم المكتبة"
    ==> "installed"
    ```

    أكثر من مكتبة:
    ------------
    ```
    lib_to_check = ["lib1", "lib2", "lib3"]
    ==>
    [
        {"lib name": "lib1", "lib status": "not found"},
        {"lib name": "lib2", "lib status": "installed"},
        {"lib name": "lib3", "lib status": "built-in"}
    ]
    ```
    """

    py_path = python_path().split("python.exe")[0] + "Lib" # C\Python312\Lib

    def check_one(lib):
        """
        اللهم صل على سيدنا محمد ﷺ
        ---------------------------

        """

        if lib in filesIn(py_path + "\\site-packages"): # تم تنزيل المكتبة
            return installed_lib_return
        elif lib in filesIn(py_path): # المكتبة أساسية في بايثون
            return builtIn_return
        else: # المكتبة غير موجودة
            return notFound_return
    
    if type(lib_to_check) == str:
        return check_one(lib_to_check)
    elif type(lib_to_check) == list:
        res = []
        for lib in lib_to_check: res.append({"lib name": lib, "lib status": check_one(lib)})
        
        return res

# file defs
def filesIn(folder_path: str = "./"):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    import os

    return os.listdir(folder_path)

def rename(file_path: str, new_name: str):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    import os

    os.rename(file_path, new_name)

def move(file_path: str, new_file_path: str = None):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------

    دالة لاقتصاص الملف/المجلد وارساله لمكان آخر\n\n

    مثال
    ------------
    ```
    file_path = "file.txt"
    new_folder_path = "./folder"
    ```
    """

    import shutil

    shutil.move(file_path, new_file_path)

def create_folder(folder_path: str):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    import os

    os.makedirs(folder_path)

def copy_folder(folder_path: str, new_folder_path: str):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    import shutil

    shutil.copytree(folder_path, new_folder_path)

def readFile(file_path: str, errorAs = 404, errorOnly: bool = False, encoding: str = "utf-8"):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    try:
        with open(file_path, 'r', encoding=encoding) as file:
            data = file.read()
            return data
    except Exception as error:
        if errorOnly == False:
            return errorAs
        else:
            return error

def saveFile(data, file_path: str, encoding: str = "utf-8", save_mode = "w"):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    with open(file_path, save_mode, encoding=encoding) as output_file:
        output_file.write(data)

def readJSON(file_path: str, errorAs = 404, errorOnly: bool = False, encoding: str = "utf-8"):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    import json

    try:
        with open(file_path, 'r', encoding=encoding) as file:
            data = json.load(file)
            return data
    except Exception as error:
        if errorOnly == False:
            return errorAs
        else:
            return error

def saveJSON(data: dict, file_path: str, indent: int = 4, encoding: str = "utf-8", save_mode = "w"):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    import json

    with open(file_path, save_mode, encoding=encoding) as output_file:
        json.dump(data, output_file, indent=indent, ensure_ascii=False)

def savePDF(merger, output_file_path, encoding = "utf-8"):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    with open(output_file_path, "wb", encoding=encoding) as output_file:
        merger.write(output_file)

def copy_file(file_path: str, new_file_path: str = None):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    import shutil

    if new_file_path is None:
        new_file_path = f"{file_path} - Copy"

    shutil.copy(file_path, new_file_path)

def delete_file(file_path):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    import os

    os.remove(file_path)

def delete_folder(folder_path):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    import shutil

    shutil.rmtree(folder_path)

# arr defs
def equalArr(arr: list, dimensions: int = 0): # arr1 = arr2
    """
    اللهم صل على سيدنا محمد ﷺ
    -----------------

    عند عمل مصفوفة واستنساخها:
    ```
    arr1 = [1, 2, 3]
    arr2 = arr1
    ```
    فإن أي تعديل تقوم به على المصفوفة الثانية يُغيّر أيضا في المصفوفة الأولى،
    ولهذا جاءت هذه الدالة: لتقوم باستنساخ المصفوفة بدون مشاكل، فقط عليك ان تدخل عدد ابعاد المصفوفة\n\n

    والفرق بين هذه الدالة وعمل: arr1.copy()\n
    هو التحكم في عدد الأبعاد التي تريد استنساخها

    ```
    
    ...
    ```

    مثلا:
    ------
    ```
    [1, 2, 3] # هذه المصفوفة لديها بُعد واحد
    [[1], [2], [3], 4] # هذه المصفوفة لديها بُعدان
    ```
    وهكذا\n
    وإذا كنت لا تعرف أبعاد المصفوفة أو تريد استنساخها بالكامل مهما كان عدد أبعادها أدخل صفر

    ```
    
    ...
    ```
    
    مثال:
    -----------
    ```
    arr1 = [1, 2, 3]
    arr2 = equalArr(arr1, 1) # أو equalArr(arr1, 0) إذا كنت تريد استنساخها مهما كان عدد أبعادها

    arr2.append(4) # إضافة رقم 4 للمصفوفة الثانية

    print(arr1) # ==> [1, 2, 3]
    print(arr2) # ==> [1, 2, 3, 4]
    ```
    """

    res = []

    if type(arr) == list:
        if dimensions > 0:
            i = 0
            while i < len(arr):
                if dimensions == 1:
                    res.append(arr[i])
                elif dimensions > 1 and type(arr[i]) == list:
                    res.append([])
                    res[i] = equalArr(arr[i], dimensions-1)
                else: # في حالة الخطأ
                    res.append(arr[i])
                
                i += 1
        else:
            res = arr.copy()
            # i = 0
            # while i < len(arr):
            #     if type(arr[i]) == list:
            #         res.append([])
            #         res[i] = equalArr(arr[i], 0)
            #     else:
            #         res.append(arr[i])
                
            #     i += 1

    return res

def allIndexInArr(arr: list, item) -> list: # 1D arr=[10,20,30,40,30], item=30 ==> [2,4]
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    ```
    arr = [10, 20, 30, 40, 30]
    item = 30
    
    ==> [2, 4]
    ```
    """

    return [index for index, value in enumerate(arr) if value == item]

def inArr(arr: list, item) -> bool: # ==> True, False
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    ==> True || False
    """

    return item in arr

def onSide(arr: list, item, greater: int = 1, startFrom: int = 0, multi_items: bool = False, want_items: bool = False): # 2D arr=[1,1,1,3,2,2,2,1,2,2], item=2, greater=1 ==> [ [4,5,6], [8,9] ] ==> [ indexArr ]
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------

    هذه الدالة تقوم بإرجاع اماكن العناصر الموضوعين بجانب بعضهم\n
    ```
    arr: المصفوفة\n
    item: العنصر الذي تريده\n
    greater: عدد العناصر التي تريد أن يكونوا بجانب بعضهم + 1 (greater: أكبر)\n
    start_from: مكان البدء (ادخل الفهرس)\n
    multi_items: هل تريد حساب أكثر من عنصر؟\n
    want_items: هل تريد ارجاع العناصر أيضا مع الفهرس؟\n\n
    ```

    مثال
    ---------
    ```
    arr=[1,1,1,3,2,2,2,1,2,2]
    item=2
    greater=1
    
    ==> [ [4,5,6], [8,9] ]
    ==> [     الفهرس     ]
    ```
    """

    finalArr = []
    onceArr = []
    finalItemsArr = []
    onceItemsArr = []
    
    i = startFrom
    while i < len(arr):
        if arr[i] == item or (type(item) == list and multi_items == True and (arr[i] in item)):
            onceArr.append(i)
            onceItemsArr.append(arr[i])
        else:
            if len(onceArr) > greater:
                finalArr.append(onceArr)
                finalItemsArr.append(onceItemsArr)

            onceArr = []
            onceItemsArr = []
        i += 1

    if len(onceArr) > greater:
        finalArr.append(onceArr)
        finalItemsArr.append(onceItemsArr)
        
        onceArr = []
        onceItemsArr = []
    

    if want_items == False:
        return finalArr
    else:
        return [finalItemsArr, finalArr]

def fillArr(arr: list, fill, fromNum: int = -1, toNum: int = -1): # arr=["⬜","⬜","⬜"], fill="🟩", from=0,to=1 ==> ["🟩","🟩","⬜"] #
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------

    ```
    arr = ["⬜","⬜","⬜"]
    fill = "🟩"
    fromNum = 0
    toNum = 1
    ```

    ==> ["🟩","🟩","⬜"]
    """

    arr = equalArr(arr)

    if fromNum == -1:
        fromNum = 0
    if toNum == -1:
        toNum = len(arr)-1

    if fromNum < len(arr) and toNum < len(arr):
        i = fromNum
        while i < toNum+1:
            arr[i] = fill
            i += 1

    return arr

def removeFromArr(arr: list, index: int = -1, txt = "", minusOne: bool = False):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    arr = equalArr(arr)
    if (index > -1) or (index <= -1 and minusOne == True):
        del arr[index%len(arr)]
    else:
        if txt in arr: arr.remove(txt)
    
    # newArr = []
    # i = 0
    # if (index > -1) or (index <= -1 and minusOne == True):
    #     while i < len(arr):
    #         if i != index%len(arr):
    #             newArr.append(arr[i])
    #         i += 1
    # else:
    #     counter = 0
    #     while i < len(arr):
    #         if arr[i] != txt or (arr[i] == txt and counter == 1):
    #             newArr.append(arr[i])
        
    #         if counter == 0 and arr[i] == txt:
    #             counter = 1
            
    #         i += 1
        
    return arr

def RemoveDuplicates(arr: list, removeOriginalItem: bool = False, specific_item: bool = False, item = None):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    def get_items(arr: list, removeOriginalItem: bool = False, specific_item: bool = False, item = None):
        duplicates = []
        unique = []

        # إضافة العناصر الأحادية والمكررة لمصفوفاتهم
        for element in arr:
            # إضافة العناصر الأحادية لمصفوفة الأحاديات
            if element not in unique or (specific_item == True and element != item): unique.append(element)

            # إضافة العناصر المكررة لمصفوفة التكرارات
            elif element not in duplicates: duplicates.append(element)

        # إلغاء العناصر (أو العنصر المحدد) الأصلي المكرر من مصفوفة الأحاديات
        if removeOriginalItem:
            for element in duplicates:
                if specific_item == False or element == item: unique = removeFromArr(unique, txt=element)
        
        return {"duplicates": duplicates, "unique": unique}
            
    data = get_items(arr, removeOriginalItem, specific_item, item)

    return data["unique"]

def sumArr(arr: list, fromNum: int = -1, toNum: int = -1): # arr=[1,2,3], fromNum=0, toNum=1 ==> 1+2 = 3
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    ```
    arr=[1, 2, 3]
    fromNum = 0
    toNum = 1
    
    ==> 1+2 = 3
    ```
    """
    
    sum = 0
    if fromNum > -1 and toNum > -1:
        toNum += 1
    elif fromNum > -1 and toNum < 0:
        toNum = len(arr)
    elif fromNum < 0 and toNum < 0:
        fromNum = 0
        toNum = len(arr)
        
    while fromNum < toNum:
        sum += arr[fromNum]
        fromNum += 1
    
    return sum

def minusArr(arr: list): # [1,4,5,8,7] ==> [3,1,3,-1]
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    ```
    arr = [1, 4, 5, 8, 7]
    
    ==> [3, 1, 3, -1]
    ```
    """

    res = []

    i = 1
    while i < len(arr):
        res.append(arr[i] - arr[i-1])
        i += 1
    
    return res

def makeColumnArr(arr: list, column: int = 0): # convert column in 2D arr to arr
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    تحويل العمود في المصفوفة ثنائية الأبعاد إلى مصفوفة منفصلة
    """
    
    arrLen = len(arr)
    columnArr = []

    i = 0
    while i < arrLen:
        if column < len(arr[i]): columnArr.append(arr[i][column])
        i += 1

    return columnArr

def convertRowToCol(arr: list, columnArr: list, column: int = 0): # Col ==> column, columnArr ==> row
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    تحويل المصفوفة أحادية الأبعاد إلى عمود في مصفوفة ثنائية الأبعاد
    """
    
    i = 0
    while i < len(columnArr):
        arr[i][column] = columnArr[i]
        i += 1
    
    return arr

def addArr(fromArr: list, toArr: list): # fromArr=[4,5,6], toArr=[1,2,3] ==> [1,2,3,4,5,6]
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    إضافة مصفوفة إلى أخرى\n

    مثال
    -------
    ```
    fromArr=[4, 5, 6]
    toArr=[1, 2, 3]
    
    ==> [1, 2, 3, 4, 5, 6]
    ```
    """
    
    return toArr + fromArr

def delArr(smallArr: list, bigArr: list): # smallArr = [4,5,6], bigArr=[1,2,3,4,5,6] ==> [1,2,3]
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    إلغاء مصفوفة من أخرى\n

    مثال
    ------------
    ```
    smallArr = [4,5,6]
    bigArr=[1,2,3,4,5,6]
    
    ==> [1,2,3]
    ```
    """
    
    smallArr = equalArr(smallArr)
    bigArr = equalArr(bigArr)

    for x in smallArr:
        if inArr(bigArr, x):
            bigArr = removeFromArr(bigArr, -1, x)
    
    return bigArr

def sliceArr(arr: list, fromNum: int = 0, toNum: int = None, step: int = 1): # arr=[10,20,30,40,50], fromNum=1, toNum=3 ==> [20,30,40]
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    ```
    arr = [10, 20, 30, 40, 50]
    fromNum = 1
    toNum = 3
    
    ==> [20, 30, 40]
    ```
    """
    
    if toNum == None: toNum = len(arr) - 1

    return arr[fromNum:toNum+1:step]
    
    # newArr = []

    # i = fromNum
    # while i <= toNum:
    #     newArr.append(arr[i])
    #     i += 1
    
    # return newArr

def reverseArr(arr: list): # arr=[1,2,3] ==> [3,2,1]
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    عكس المصفوفة

    مثال
    ---------
    ```
    arr=[1, 2, 3]
    
    ==> [3, 2, 1]
    ```
    """

    return arr[::-1]

def search2D(arr: list, item, column: int = 0): # ==> [index]
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    البحث عن عنصر في مصفوفة ثنائية الأبعاد عن طريق العمود
    """

    res = []
    ind = []

    i = 0
    while i < len(arr):
        if arr[i][column] == item:
            res.append(arr[i])
            ind.append(i)
        i += 1

    return [res, ind]

def convert2D(arr: list): # ==> 1D, convert 2D to 1D, arr=[ [1,2,3], [4,5,6], [7,8,9] ] ==> [1,2,3,4,5,6,7,8,9]
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    تحويل المصفوفة ثنائية الأبعاد إلى أحادية الأبعاد\n

    مثال
    -----------
    ```
    arr=[ [1,2,3], [4,5,6], [7,8,9] ]
    
    ==> [1,2,3,4,5,6,7,8,9]
    ```
    """

    newArr = []
    for item in arr:
        if type(item) == list: newArr += item
        else: newArr.append(item)

    return newArr

def rotate2DArr(arr: list, numR: int = 1): # numR ==> num of right turns
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    arr = equalArr(arr, 2)
    
    a = 0
    while a < numR:
        res = []

        i = 0
        while i < len(arr[0]):
            res.append(makeColumnArr(arr, i)[::-1])
            i += 1
        
        arr = res
        a += 1
    
    return arr

def insertInArr(arr: list, itemToInsert, itemIndex: int = 0, multi_items: bool = False): # arr=[1,3,4], itemToInsert=2 ==> [1,2,3,4]
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    ```
    arr=[1,3,4]
    itemToInsert = 2
    itemIndex = 1

    ==> [1,2,3,4]
    ```
    """

    res = equalArr(arr)
    if multi_items and type(itemToInsert) == list:
        for i, item in enumerate(itemToInsert): res.insert(itemIndex, itemToInsert[len(itemToInsert)-i-1])
    else:
        res.insert(itemIndex, itemToInsert)
    
    # arr = equalArr(arr)
    # res = []

    # i = 0
    # while i < len(arr):
    #     if i == itemIndex:
    #         if multi_items == False: res.append(itemToInsert)
    #         elif type(itemToInsert) == list:
    #             e = 0
    #             while e < len(itemToInsert):
    #                 res.append(itemToInsert[e])
    #                 e += 1
        
    #     res.append(arr[i])
    #     i += 1

    return res

def convertTxt(text: str): # convert text to code, "5 + 7" ==> 12
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    تحويل النص لشيفرة\n

    مثال
    -----------------
    ```
    "5 + 7" ==> 12
    ```
    """
    
    try:
        return eval(text)
    except Exception as e:
        print("حدث خطأ في التحويل:", e)
        return None

def searchIn_decodeNum(arr: list, search_key, equalOrContain: int = 0): # arr=["n1", "o1", "n2"], search_key="n" ==> [0, 2] => [index]
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    ```
    arr=["n1", "o1", "n2"]
    search_key = "n"
    
    ==> [0, 2]
    ==> [فهرس]
    ```
    """
    
    arr = equalArr(arr)
    result = []
    
    newArr = []
    i = 0
    while i < len(arr):
        newArr.append(decodeNum(arr[i]))
        for index, item in enumerate(newArr[i]):
            newArr[i][index] = str(item)

        if (equalOrContain == 0 and search_key in newArr[i]) or (equalOrContain == 1 and search_key in "".join(newArr[i])):
            result.append(i)
        i += 1

    return result

def searchInObjArr(arr: list, key_want_search: str, value_want_search, regex_value: bool = False, multi_values: bool = False): # [ {"id": 1, "num": 453}, {"id": 2, "num": 734} ], id == 1 ==> [  [ {"id": 1, "num": 453} ], [0]  ]
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    ```
    [ {"id": 1, "num": 453}, {"id": 2, "num": 734} ]
    key_want_search = "id"
    value_want_search = 1
    
    ==> [  [ {"id": 1, "num": 453} ], [0]  ]
    ```
    """

    import re

    res = []
    ind = []
    for i, item in enumerate(arr):
        if type(item) == dict:
            if key_want_search in item:
                if ((multi_values == False and                                     (item[key_want_search] == value_want_search or (regex_value == True and re.match(value_want_search, item[key_want_search])))) or
                    (multi_values == True  and type(value_want_search) == list and (item[key_want_search] in value_want_search or (regex_value == True and      any(re.search(pattern, item[key_want_search]) for pattern in value_want_search))))):
                    res.append(item)
                    ind.append(i)
    
    return [res, ind]

# object {} defs
def equalObject(the_object: dict, dimensions: int = 0): # the_object1 = the_object2
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    
    عند عمل قاموس واستنساخه:
    ```
    obj1 = {"number": 1}
    obj2 = obj1
    ```
    فإن أي تعديل تقوم به على القاموس الثاني يُغيّر أيضا في القاموس الأول،
    ولهذا جاءت هذه الدالة: لتقوم باستنساخ القاموس بدون مشاكل، فقط عليك ان تدخل عدد ابعاد القاموس\n\n

    والفرق بين هذه الدالة وعمل: obj1.copy()\n
    هو التحكم في عدد الأبعاد التي تريد استنساخها

    ```
    
    ...
    ```

    مثلا:
    ------
    ```
    {"number": 1} # هذا القاموس لديه بُعد واحد
    {"number": {"number2": 2}} # هذا القاموس لديه بُعدان
    ```
    وهكذا\n
    وإذا كنت لا تعرف أبعاد القاموس أو تريد استنساخه بالكامل مهما كان عدد أبعاده أدخل صفر

    ```
    
    ...
    ```
    
    مثال:
    -----------
    ```
    obj1 = {"numbers":  {"n1": 1, "n2": 2, "n3": 3}  } # القاموس الأول لديه بُعدان
    obj2 = equalObject(obj1, 2) # أو equalObject(obj1, 0) إذا كنت تريد استنساخه مهما كان عدد أبعاده

    obj2["numbers"]["n4"] = 4 # إضافة رقم 4 للقاموس الثاني

    print(obj1) # ==> {"numbers": {"n1": 1, "n2": 2, "n3": 3}}
    print(obj2) # ==> {"numbers": {"n1": 1, "n2": 2, "n3": 3, 4}}
    ```
    """
    
    res = {}
    if type(the_object) == dict:
        if dimensions > 0:
            for key in the_object:
                item = the_object[key]
                if dimensions == 1:
                    res[key] = item
                elif dimensions > 1 and type(item) == dict:
                    res[key] = equalObject(item, dimensions-1)
                else: # في حالة الخطأ
                    res[key] = item
        else:
            res = the_object.copy()
            # for key in the_object:
            #     item = the_object[key]
            #     if type(item) == dict:
            #         res[key] = equalObject(item, 0)
            #     else:
            #         res[key] = item                

    return res

def getObjectKeys(the_object: dict): # {"h1": "text", "h2": "text2"} ==> ["h1", "h2"]
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    ```
    the_object = {"h1": "text", "h2": "text2"}
    
    ==> ["h1", "h2"]
    ```
    """

    return list(the_object.keys())

def addToObject(the_object: dict, key: str, want_to_add): # {"name": "about"}, "name", "about" ==> {"name": "about", "name 2": "about"}
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    ```
    the_object = {"name": "about"}
    key = "name"
    value = "about"
    
    ==> {"name": "about", "name 2": "about"}
    ```
    """

    the_object = equalObject(the_object)
    if len(searchIn_decodeNum(getObjectKeys(the_object), key, 1)) > 0:
        the_object[f"{key} {len(searchIn_decodeNum(getObjectKeys(the_object), key, 1))+1}"] = want_to_add
    else:
        the_object[key] = want_to_add

    return the_object

def key_index(the_object: dict, key: str):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    مثال
    -----------
    ```
    the_object = {"number 1": 10, "number 2": 20, "number 3": 30}
    key = "number 3"

    ==> 2
    ```
    
    مثال 2
    -----------
    ```
    the_object = {"number 1": 10, "number 2": 20, "number 3": 30}
    key = "number 5"

    ==> -1
    ```
    """
    
    if key in the_object:
        return getObjectKeys(the_object).index(key)
    else:
        return -1

def delete_key(the_object: dict, key: str): # {"n1":10, "n2":20, "n3":30}, key = "n1" ==> {"n2":20, "n3":30}
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    مثال
    ----------
    ```
    the_object = {"number 1": 10, "number 2": 20, "number 3": 30}
    key = "number 1"
    ==> {"number 2": 20, "number 3": 30}\n\n
    ```

    مثال 2
    ----------
    ```
    the_object = {"number 1": 10, "number 2": 20, "number 3": 30}
    key = ["number 1", "number 2"]
    ==> {"number 3": 30}
    ```
    """

    the_object = equalObject(the_object)

    if type(key) == str:
        if key in the_object: the_object.pop(key)
    elif type(key) == list:
        for item in key:
            the_object = delete_key(the_object, item)

    return the_object

def insert_key(the_object: dict, i: int, key: str, value = None, change_value: bool = False):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    مثال
    -----------
    ```
    the_object = {"num1": 10, "num3": 30}
    i = 1
    key = "num2"
    value = 20

    ==> {"num1": 10, "num2": 20, "num3": 30}
    ```

    مثال 2
    -----------
    ```
    the_object = {"num1": 10, "num3": 30, "num2": 20}
    i = 1
    key = "num2"

    ==> {"num1": 10, "num2": 20, "num3": 30}
    ```

    مثال 3
    -----------
    ```
    the_object = {"num1": 10, "num3": 30, "num2": 50}
    i = 1
    key = "num2"
    value = 20
    change_value = True

    ==> {"num1": 10, "num2": 20, "num3": 30}
    ```
    """

    obj_len = len(getObjectKeys(the_object))

    tmp_list = list(the_object.items())
    if key in the_object:
        tmp_value = the_object[key]
        tmp_list.remove((key, the_object[key]))

        if change_value == True:
            return {**dict(tmp_list[:(i%obj_len)]), key: value, **dict(tmp_list[(i%obj_len):])}
        else:
            return {**dict(tmp_list[:(i%obj_len)]), key: tmp_value, **dict(tmp_list[(i%obj_len):])}
    else:
        return {**dict(tmp_list[:(i%obj_len)]), key: value, **dict(tmp_list[(i%obj_len):])}

def change_key(the_object: dict, old_key: str, new_key: str): # {"number 1": 10, "number 5": 20, "number 3": 30}, "number 5", "number 2" ==> {"number 1": 10, "number 2": 20, "number 3": 30} #
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    ```
    the_object = {"number 1": 10, "number 5": 20, "number 3": 30}
    old_key = "number 5"
    new_key = "number 2"

    ==> {"number 1": 10, "number 2": 20, "number 3": 30}
    ```
    """

    the_object = equalObject(the_object)

    if old_key in the_object and new_key not in the_object:
        key_ind = key_index(the_object, old_key)
        value = the_object[old_key]

        the_object = delete_key(the_object, old_key)
        the_object = insert_key(the_object, key_ind, new_key, value)
    
    return the_object

# json defs
def combine_json_files(files_names: list, output_file_path: str, indent: int = 4, encoding: str = "utf-8"):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    دمج ملفات json
    """
    
    combined_data = {}

    # قراءة محتوى كل ملف وإضافة المعلومات إلى القاموس
    for file_name in files_names:
        data = readFile(file_name, encoding)
        combined_data.update(data)

    # حفظ الملف
    saveJSON(combined_data, output_file_path, indent, encoding)

def uncombine_json_file(combined_file_path: str, indent: int = 4, encoding: str = "utf-8", files: str = 'f"{key}.json"'):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    فصل ملفات json
    """

    # قراءة الملف
    combined_data = readFile(combined_file_path, encoding)

    # حفظ الملفات
    for key in combined_data:
        saveJSON({key:combined_data[key]}, eval(files), indent, encoding)

# pdf defs
def combine_pdf_files(files_names: list, output_file_path: str, encoding: str = "utf-8"):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    دمج ملفات القراءة pdf
    """
    
    import PyPDF2

    merger = PyPDF2.PdfMerger()

    for pdf in files_names:
        merger.append(pdf)

    savePDF(merger, output_file_path, encoding)

def pdf_page_to_elements(pdf_path: str, page_num: int) -> list[dict]:
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    دالة ترجع عناصر صفحة واحدة ملف القراءة pdf
    """

    import fitz  # PyMuPDF
    import base64

    def convert_bytes(data):
        if isinstance(data, bytes):
            try:
                return data.decode('utf-8')
            except UnicodeDecodeError:
                # تحويل البيانات لـ base64 إذا فشل التحويل لـ utf-8
                return base64.b64encode(data).decode('utf-8')
        if isinstance(data, dict):
            return {key: convert_bytes(value) for key, value in data.items()}
        if isinstance(data, list):
            return [convert_bytes(element) for element in data]
    
        return data
    
    def dict_tuple(obj: dict):
        obj = obj.copy()

        for key in obj:
            item = obj[key]
            if type(item) == list: obj[key] = tuple_to_list(item)
            elif type(item) == dict: obj[key] = dict_tuple(item)
            elif type(item) == tuple:
                obj[key] = list(obj[key])
                obj[key] = tuple_to_list(obj[key])

        return obj

    def tuple_to_list(arr: list):
        arr = arr.copy()

        for i, item in enumerate(arr):
            if type(item) == list: arr[i] = tuple_to_list(item)
            elif type(item) == dict: arr[i] = dict_tuple(item)
            elif type(item) == tuple:
                arr[i] = list(arr[i])
                arr[i] = tuple_to_list(arr[i])
        
        return arr
    
    doc = fitz.open(pdf_path)
    
    page = doc.load_page(page_num)
    blocks = page.get_text("dict")["blocks"]
    
    return convert_bytes(tuple_to_list(blocks))

def pdf_to_elements(pdf_path: str, from_page: int = 1, to_page: int = -1) -> dict[list]:
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    دالة ترجع عناصر ملف القراءة pdf
    """

    import fitz  # PyMuPDF
    import base64

    def convert_bytes(data):
        if isinstance(data, bytes):
            try:
                return data.decode('utf-8')
            except UnicodeDecodeError:
                # تحويل البيانات لـ base64 إذا فشل التحويل لـ utf-8
                return base64.b64encode(data).decode('utf-8')
        if isinstance(data, dict):
            return {key: convert_bytes(value) for key, value in data.items()}
        if isinstance(data, list):
            return [convert_bytes(element) for element in data]
    
        return data
    
    def dict_tuple(obj: dict):
        obj = obj.copy()

        for key in obj:
            item = obj[key]
            if type(item) == list: obj[key] = tuple_to_list(item)
            elif type(item) == dict: obj[key] = dict_tuple(item)
            elif type(item) == tuple:
                obj[key] = list(obj[key])
                obj[key] = tuple_to_list(obj[key])

        return obj

    def tuple_to_list(arr: list):
        arr = arr.copy()

        for i, item in enumerate(arr):
            if type(item) == list: arr[i] = tuple_to_list(item)
            elif type(item) == dict: arr[i] = dict_tuple(item)
            elif type(item) == tuple:
                arr[i] = list(arr[i])
                arr[i] = tuple_to_list(arr[i])
        
        return arr

    if to_page < 0: to_page = pdf_pages_num(pdf_path)
    
    res = {}
    doc = fitz.open(pdf_path)
    
    i = from_page
    while i <= to_page:
        page = doc.load_page(i)
        blocks = page.get_text("dict")["blocks"]
        res[str(i)] = blocks
        i += 1
    
    return convert_bytes(dict_tuple(res))

def pdf_pages_num(pdf_path: str) -> int:
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    دالة ترجع عدد صفحات ملف القراءة pdf
    """

    import fitz  # PyMuPDF
    doc = fitz.open(pdf_path)

    return len(doc)

def images_to_pdf(image_paths: list[str], output_pdf: str):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    دمج الصور في ملف pdf
    """

    from PIL import Image
    from fpdf import FPDF
    import os

    pdf = FPDF()

    for image_path in image_paths:
        image = Image.open(image_path)
        
        # image to RGB (fpdf required)
        if image.mode in ("RGBA", "LA"):
            background = Image.new("RGB", image.size, (255, 255, 255))
            background.paste(image, mask=image.split()[3])  # Handle transparency
            image = background
        
        # make image
        base_name = os.path.splitext(image_path)[0]  # Get the file name without extension
        image_path_jpg = f"{base_name}.jpg"  # Create a new file name with .jpg extension
        image.save(image_path_jpg, "JPEG")  # Save in JPEG format
        
        # put image on pdf
        pdf.add_page()
        pdf.image(image_path_jpg, 0, 0, pdf.w, pdf.h)
    
    pdf.output(output_pdf)

# xlsx defs
def xlsx_to_arr(xlsx_path: str) -> list[list]:
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    تحويل ملف xlsx إلى مصفوفة ثنائية الأبعاد
    """

    from openpyxl import load_workbook

    # افتح الملف
    workbook = load_workbook(xlsx_path)

    # اختر الورقة التي تحتوي على البيانات
    sheet = workbook.active

    # انشئ مصفوفة لتخزين البيانات
    res = []

    for i, row in enumerate(sheet.iter_rows(values_only=True)):
        row = list(row)
        res.append(row)

    return res

def arr_to_xlsx(arr: list, xlsx_path: str):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    import openpyxl

    # إنشاء ملف إكسل جديد
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    
    # تعبئة البيانات في الورقة
    for i, row in enumerate(arr):
        for j, value in enumerate(row):
            sheet.cell(row=i+1, column=j+1, value=value)
    
    # حفظ الملف
    workbook.save(xlsx_path)

# docx defs
def docx_to_text(docx_path: str):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    from docx import Document

    # فتح مستند الوورد
    doc = Document(docx_path)
    
    # قراءة محتوى المستند
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    
    # دمج النصوص في نص واحد
    return '\n'.join(full_text)

def text_to_docx(text: str, docx_path: str):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """
    
    from docx import Document
    
    # إنشاء مستند وورد جديد
    doc = Document()
    
    # إضافة النص للمستند
    doc.add_paragraph(text)
    
    # حفظ الملف
    doc.save(docx_path)

# text defs
def decodeNum(text: str): # r2 ==> ["r", 2]
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    فصل الأرقام عن النصوص

    ```
    data_input = "r2"
    
    ==> ["r", 2]
    ```
    """

    res = []
    tmp = ""
    nums = "0123456789"

    for i in rangeNum(len(text)-1):
        tmp += text[i]
        if text[i] not in nums:
            if text[i+1] in nums:
                res.append(tmp)
                tmp = ""
        else:
            if text[i+1] not in nums:
                res.append(int(tmp))
                tmp = ""
    
    tmp += text[-1]
    if text[-1] in nums: res.append(int(tmp))
    if text[-1] not in nums: res.append(tmp)
    tmp = ""

    return res

def right(text: str, amount: int):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    return text[-amount:]

def left(text: str, amount: int):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    return text[:amount]

def mid(text: str, offset: int, amount: int):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    return text[offset:offset+amount]

def split(text: str, splitter: str | list[str] = ""):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    هذه الدالة تقسم النص
    """
    
    res = []
    if splitter == "":
        for char in text:
            res.append(char)
    else:
        if type(splitter) == str:
            res = text.split(splitter)
        elif type(splitter) == list:
            if splitter == []:
                res = text
            else:
                for splitter_item in splitter:
                    if res == []:
                        res.append(split(text, splitter_item))
                    else:
                        res_len = len(res)
                        res_i = res_len-1
                        while res_i >= 0:
                            res = convert2D(res)
                            res_item = res[res_i]

                            res[res_i] = split(res_item, splitter_item)
                            res = convert2D(res)

                            res_i -= 1

    return res

def removeSpaces(text: str, strip_text: bool = False): # إلغاء   المسافات
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    إلغاء   المسافات
    """
    
    text = " ".join(RemoveDuplicates(text.split(" "), True, True, ""))
    text = " ".join(removeFromArr(text.split(" "), txt=""))

    return text.strip() if strip_text else text

def replaceText(text: str, replArr: list[list[str]]): # text = "hello word!", replArr = [ ["word", "world"], ["!", "."] ] ==> "hello world."
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    ```
    text = "hello word!"
    replArr = [ ["word", "world"], ["!", "."] ]
    ```
    \n\n
    ==> "hello world."
    """
    
    i = 0
    while i < len(replArr):
        text = text.replace(replArr[i][0], replArr[i][1])
        i += 1

    return text

# num defs
def makeZeroNum(num: str | int, cells_num: int = 3, zero: str = "0"):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------

    ```
    num = "10"
    cells_num = 4
    ==> "0010"
    ```
    """

    if zero == "0":
        return eval('f"{' + str(num) + ':0' + str(cells_num) + '}"')
    else:
        num = str(num)
        if cells_num > len(num):
            for item in rangeNum(cells_num - len(num)):
                num = zero + num
    
        return num

def toTime(seconds: int):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------

    تحويل الثواني لوقت مقسم

    ```
    640 ==> "00:10:40"
    ```
    """
    
    hours = seconds // 3600
    remaining_seconds = seconds % 3600
    minutes = remaining_seconds // 60
    remaining_seconds = remaining_seconds % 60
    
    return f"{hours:02}:{minutes:02}:{remaining_seconds:02}"

# separate defs
def separate_text(text: str, splitter: list = ["."], joiner: str = "\n", stripping: bool = False, delete_splitter: bool = False): # "line1.line2" ==> "line1.\nline2"
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    
    ```
    text = "line1.line2"
    splitter = ["."] # بعد النقطة ضع سطرًا جديدًا
    joiner = "\\n"

    ==> "line1.\\nline2"
    ```
    """
    
    if type(splitter) == list:
        for item in splitter:
            tmp = text.split(item)
            if stripping == True:
                for tmp_i, tmpItem in enumerate(tmp): tmp[tmp_i] = tmpItem.strip()
            
            if delete_splitter == False: text = f"{item}{joiner}".join(tmp)
            else: text = joiner.join(tmp)
    else:
        tmp = text.split(str(splitter))
        if stripping == True:
            for tmp_i, tmpItem in enumerate(tmp): tmp[tmp_i] = tmpItem.strip()
        
        if delete_splitter == False: text = f"{splitter}{joiner}".join(tmp)
        else: text = joiner.join(tmp)
    
    return text

def separate_arr(arr: list, splitter: list = ["."], joiner: str = "\n", stripping: bool = False, delete_splitter: bool = False):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    new_arr = []

    for item in arr:
        if type(item) == str:
            new_arr.append(separate_text(item, splitter, joiner, stripping, delete_splitter))
        elif type(item) == list:
            new_arr.append(separate_arr(item, splitter, joiner, stripping))
        elif type(item) == dict:
            new_arr.append(separate_obj(item, splitter, joiner, stripping))
        else:
            new_arr.append(item)
    
    return new_arr

def separate_obj(obj: dict, splitter: list = ["."], joiner: str = "\n", stripping: bool = False, delete_splitter: bool = False):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    new_obj = {}

    for key in obj:
        item = obj[key]
        if type(item) == str:
            new_obj[key] = separate_text(item, splitter, joiner, stripping, delete_splitter)
        elif type(item) == list:
            new_obj[key] = separate_arr(item, splitter, joiner, stripping)
        elif type(item) == dict:
            new_obj[key] = separate_obj(item, splitter, joiner, stripping)
        else:
            new_obj[key] = item
    
    return new_obj

# translate defs
def translate_text(text, target_language: str ="en", not_langs: list = [], detectOrSrc: int = 0):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    from googletrans import Translator
    from langdetect import detect
    
    translator = Translator()

    if not_langs == []:
        if text:
            textArr = []
            for i in text.split("\n"):
                if i != "":
                    tmp = ""
                    try:
                        tmp = translator.translate(i, dest=target_language)
                        textArr.append(tmp.text)
                    except:
                        tmp = i
                        textArr.append(tmp)

            translation = "\n".join(textArr)
            return translation
    else:
        sentences = text.split(" ")
    
        fullArr = []
        fullArrInd = []
        tmpArr = []
        tmpInd = -1
        # ["hello", "اهلا"] ==> fullArr = [ ["hello"] ], fullArrInd = [ 0 ]
        detected_language = ""
        for i, item in enumerate(sentences):
            detected_language = detect_lang(item, detectOrSrc)
            if detected_language != 404:
                if detected_language not in not_langs:
                    tmpArr.append(item)
                    if tmpInd == -1: tmpInd = i
                elif i != 0 and len(tmpArr) > 0:
                    fullArr.append(tmpArr)
                    fullArrInd.append(tmpInd)
                    tmpArr = []
                    tmpInd = -1
            else:
                fullArr.append(item)
                fullArrInd.append(i)
                tmpArr = []
                tmpInd = -1
                detected_language = ""
        if detected_language not in not_langs and len(tmpArr) > 0:
            fullArr.append(tmpArr)
            fullArrInd.append(tmpInd)
            tmpArr = []
            tmpInd = -1
        
        for i, item in enumerate(fullArr):
            i = len(fullArr)-i-1
            for e in range(len(fullArr[i])):
                sentences = removeFromArr(sentences, fullArrInd[i])
            sentences.insert(fullArrInd[i], fullArr[i])

        for i, item in enumerate(sentences):
            if type(item) == list:
                tmpTransText = " ".join(item)
                tmpTrans = ""
                try: tmpTrans = translator.translate(tmpTransText, dest=target_language).text
                except: tmpTrans = tmpTransText
                sentences[i] = tmpTrans

        text = " ".join(sentences)
    
    return text

def translate_arr(arr, target_language: str ="en", not_langs: list = []):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    newArr = []
    for item in arr:
        if type(item) == str:
            newArr.append(translate_text(item, target_language, not_langs))
        elif type(item) == dict: # dict ==> {}
            newArr.append(translate_obj(item, target_language, not_langs))
        elif type(item) == list: # list ==> []
            newArr.append(translate_arr(item, target_language, not_langs))
        else:
            newArr.append(item)
        
    return newArr

def translate_obj(obj: dict, target_language: str = "en", not_langs: list = []):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    new_obj = {}

    for i in obj:
        if type(obj[i]) == str:
            new_obj[translate_text(i, target_language, not_langs)] = translate_text(obj[i], target_language, not_langs)
        elif type(obj[i]) == dict: # dict ==> {}
            new_obj[translate_text(i, target_language, not_langs)] = translate_obj(obj[i], target_language, not_langs)
        elif type(obj[i]) == list: # list ==> []
            new_obj[translate_text(i, target_language, not_langs)] = translate_arr(obj[i], target_language, not_langs)
        else:
            new_obj[translate_text(i, target_language, not_langs)] = obj[i]

    return new_obj

# decode html
def decode_html_in_text(text: str, tag: str = "img", attribute: str = "src", endTag: bool = False, delete_inner_html: bool = False):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    
    مثال:
    --------------
    ```
    text = '<img src="the_link.com">'
    tag = "img"
    attribute = "src"
    endTag = 0
    delete_inner_html = 0 | 1
    ```

    ...\n
    النتيجة:
    ---
    the_link.com
    """

    import re

    if attribute != "":
        check = rf'<{tag} .*?{attribute}="([^"]*)".*?>'
        if endTag == True: check += rf'.*?</{tag}>'
    else: # inner html
        check = rf'<{tag}.*?>'
        if delete_inner_html == False:
            check += rf'([^"]*)'
        elif endTag == True:
            check += rf'.*?'
        if endTag == True: check += rf'</{tag}>'

    return re.sub(check, r"\1", text)

def decode_html_in_arr(arr: list, tag: str = "img", attribute: str = "src", endTag: bool = False, delete_inner_html: bool = False):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    newArr = []
    for item in arr:
        if type(item) == list:
            newArr.append(decode_html_in_arr(item, tag, attribute, endTag, delete_inner_html))
        elif type(item) == dict:
            newArr.append(decode_html_in_obj(item, tag, attribute, endTag, delete_inner_html))
        elif type(item) == str:
            newArr.append(decode_html_in_text(item, tag, attribute, endTag, delete_inner_html))
        else:
            newArr.append(item)

    return newArr

def decode_html_in_obj(obj: dict, tag: str = "img", attribute: str = "src", endTag: bool = False, delete_inner_html: bool = False):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    newObj = {}
    for key in obj:
        item = obj[key]
        if type(item) == list:
            newObj[key] = decode_html_in_arr(item, tag, attribute, endTag, delete_inner_html)
        elif type(item) == dict:
            newObj[key] = decode_html_in_obj(item, tag, attribute, endTag, delete_inner_html)
        elif type(item) == str:
            newObj[key] = decode_html_in_text(item, tag, attribute, endTag, delete_inner_html)
        else:
            newObj[key] = item
    
    return newObj

# keyboard defs
def write(text: str, delay: float = 0, restore_state_after: bool = True, exact: bool | None = None):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    import keyboard

    keyboard.write(text, delay, restore_state_after, exact)

def press(hot_key):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    import keyboard

    keyboard.press(hot_key)

def unpress(hot_key):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    import keyboard

    keyboard.release(hot_key)

def copy(text: str):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------

    نسخ النص للحافظة
    """

    import pyperclip

    pyperclip.copy(text)

# mouse defs
def getMousePosition() -> list[int]:
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    import mouse

    return [mouse.get_position()[0], mouse.get_position()[1]]

def getMousePositions(saveKey: str ="g", breakKey: str ="0"):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    import keyboard

    all = []

    while True:
        if keyboard.is_pressed(saveKey):
            mouse_position = getMousePosition()
            item = input("what is it?\n")
            all.append([ item, mouse_position ])

        if keyboard.is_pressed(breakKey):
            break

    return all

def click():
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    import mouse

    mouse.click()

def right_click():
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    import mouse

    mouse.right_click()

def move(x: int, y: int, absolute: bool = True, duration: int = 0):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    import mouse

    mouse.move(x, y, absolute, duration)

# chatGPT
def chatGPT(user_question: str, api_key: str):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    import openai

    openai.api_key = api_key

    response = openai.completions.create(
        model = "gpt-3.5-turbo",
        messages = [
            {
                "role": "user",
                "content": user_question
            }
        ]
    )

    answer = response["choices"][0]["message"]["content"]
    return answer

# soup
def get_soup(link: str, wait: int = 0):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    from bs4 import BeautifulSoup
    import requests
    import time

    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.3'
    }

    with requests.Session() as session:
        try:
            result = session.get(link, headers=headers, timeout=10)
            time.sleep(wait)
            result.raise_for_status()  # رفع استثناء إذا كان هناك خطأ في الاستجابة
            src = result.content
            soup = BeautifulSoup(src, "lxml")
            return soup
        except requests.exceptions.RequestException as e:
            # print(f"Error: {e}")
            return None

def get_file_soup(file_path: str):
    from bs4 import BeautifulSoup

    return BeautifulSoup(readFile(file_path), "html.parser")

def get_redirected_url(link: str): # bit.ly ==> link
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    import requests

    # إجراء طلب GET للحصول على الصفحة الأولى
    response = requests.get(link)

    # التحقق من وجود عملية تحويل
    if response.history:
        # الحصول على الرابط النهائي بعد التحويل
        final_url = response.url
        return final_url
    else:
        return link

def innerText(item):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    return item.text

def innerHTML(item):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    return "".join(map(str, item.contents))

def delete_text_out_of_html_elements(item):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------

    حذف النص خارج العناصر

    ```
    soup = \"\"\"<div>
        نَصّ خارج العناصر
        <h1>نَصّ داخل عنصر</h1>
    </div>\"\"\"

    element = soup.select(\"div\")
    element = delete_text_out_of_html_elements(element)
    ```

    الناتج:
    ```
    <div>
        <h1>نَصّ داخل عنصر</h1>
    </div>
    ```
    """

    from bs4 import NavigableString
    
    for content in item.contents:
        if isinstance(content, NavigableString):
            content.extract() # حذف النص الخارجي

    return item

# selenium
def makeFirefoxDriver():
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    from selenium import webdriver
    from selenium.webdriver.common.by import By

    options = webdriver.FirefoxOptions()
    options.add_argument('--headless')
    driver = webdriver.Firefox()

    return driver

def driverToSoup(driver):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    from bs4 import BeautifulSoup

    return BeautifulSoup(driver.page_source, "html.parser")

# dataURL
def imageLink_to_dataurl(image_url, extension = "png", error = 404):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    import base64
    import requests

    # تنزيل الصورة من الرابط
    response = requests.get(image_url)
    
    # التحقق من أن الاستجابة صحيحة
    if response.status_code == 200:
        data_url = f'data:image/{extension};base64,' + base64.b64encode(response.content).decode('utf-8')
        return data_url
    else:
        return error

def imageFile_to_dataurl(image_path, extension="png", error=404):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    import base64

    try:
        # قراءة محتوى الملف كبايتس
        with open(image_path, "rb") as file:
            image_data = file.read()

        # التحويل إلى رابط data URL
        data_url = f"data:image/{extension};base64," + base64.b64encode(image_data).decode("utf-8")
        return data_url
    except FileNotFoundError:
        return error

def videoFile_to_dataurl(video_path, error = 404):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    import base64

    try:
        # قراءة ملف الفيديو كتسلسل بيانات
        with open(video_path, "rb") as video_file:
            video_data = video_file.read()

        # تحويل التسلسل إلى base64
        video_base64 = base64.b64encode(video_data)

        # إنشاء Data URL
        data_url = "data:video/mp4;base64," + video_base64.decode('utf-8')

        return data_url
    except:
        return error

# youtube
def video_trans(video_id: str, languages: list = ["ar"]) -> list[dict]:
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------

    دالة لإرجاع ترجمة المقطع على يوتيوب

    ```
    ==> [ {"text": "أهلا بكم", "duration": 1.2, "start": 0} ]
    ```
    """

    from youtube_transcript_api import YouTubeTranscriptApi
    
    return YouTubeTranscriptApi.list_transcripts(video_id).find_transcript(languages).fetch()

# download
def download_file(file_url: str, output_file_path: str, show_progress_bar: bool = True, print_done: bool = True):
    import requests
    import urllib.parse
    from tqdm import tqdm

    response = requests.get(file_url, stream=True)
    if response.status_code == 200:
        total_size = int(response.headers.get('content-length', 0))
        block_size = 1024  # 1 Kibibyte
        if show_progress_bar: progress_bar = tqdm(total=total_size, unit='iB', unit_scale=True)

        with open(output_file_path, 'wb') as f:
            for data in response.iter_content(block_size):
                if show_progress_bar: progress_bar.update(len(data))
                f.write(data)

        if show_progress_bar: progress_bar.close()
        if print_done: print(f"done: {urllib.parse.unquote(file_url, encoding='utf-8')}")
    else:
        if print_done: print(f"error: {urllib.parse.unquote(file_url, encoding='utf-8')}.\nerror code: {response.status_code}")

# done message
def doneMessage(want_clear: int = 1):
    """
    اللهم صل على سيدنا محمد ﷺ
    ---------------------------
    """

    if want_clear > 0: clear()
    print("Done.. Alhamdulillah")