import os
import pypdf
import zipfile
import tempfile
import docx.api
import spire.doc
from docx import Document
from arm_rag.config import CONFIG

import importlib
import gmft
import gmft.table_detection
import gmft.table_visualization
import gmft.table_function
import gmft.table_function_algorithm
import gmft.table_captioning
import gmft.pdf_bindings.bindings_pdfium
import gmft.pdf_bindings
import gmft.common

import pandas as pd
import json
import re

from gmft.pdf_bindings import PyPDFium2Document
from gmft.auto import CroppedTable, AutoTableDetector
from gmft.auto import AutoTableFormatter

importlib.reload(gmft)
importlib.reload(gmft.common)
importlib.reload(gmft.table_captioning)
importlib.reload(gmft.table_detection)
importlib.reload(gmft.table_visualization)
importlib.reload(gmft.table_function)
importlib.reload(gmft.table_function_algorithm)
importlib.reload(gmft.pdf_bindings.bindings_pdfium)
importlib.reload(gmft.pdf_bindings)

class Document_parser:
    def __init__(self):
        self.supported_formats = CONFIG['document']['supported_formats']
    

    def parse(self, path):
        """The main method which parses initial file/archive  from the user"""
        
        if path.endswith('.zip'):
            return self.parse_archive(path) # returns dict of multiple documents
        else:
            content, tables = self.parse_file(path)
            _, filename = os.path.split(path)
            return {filename: content}, {f'{filename}_table': tables} # returns dict of signle document


    def parse_archive(self, archive_path):
        """returns a dict: {filename: content}"""
        file_contents = {}
        with zipfile.ZipFile(archive_path, 'r') as zip_ref:
            with tempfile.TemporaryDirectory() as temp_dir:
                for member in zip_ref.namelist():
                    extracted_path = zip_ref.extract(member, path=temp_dir)
                    content = self.parse_file(extracted_path)
                    file_contents[member] = content
        return file_contents


    def parse_file(self, file_path):
        if '.' + file_path.split('.')[-1] in self.supported_formats:
            if file_path.endswith('.pdf'):
                content = self.read_pdf(file_path)
                tables = self.read_tables_pdf(file_path)
            elif file_path.endswith('.txt'):
                content = self.read_txt(file_path)
                tables = None
            elif file_path.endswith('.doc'):
                content = self.read_doc(file_path)
                tables = None
            elif file_path.endswith('.docx'):
                content = self.read_docx(file_path)
                tables = self.read_tables_docx(file_path)
            return content, tables
        else:
            raise ValueError(f"The file type is not supported. Please upload only {str(self.supported_formats)[1:-1]}.")
    
    def read_tables_pdf(self, path):
        """Read tables from a pdf file"""
        def make_unique_columns(columns):
            seen = {}
            new_columns = []
            for item in columns:
                if item in seen:
                    seen[item] += 1
                    new_columns.append(f"{item}_{seen[item]}")
                else:
                    seen[item] = 0
                    new_columns.append(item)
            return new_columns
        
        def extract_tables(pdf_path) -> list[CroppedTable]:
            detector = AutoTableDetector()
            doc = PyPDFium2Document(pdf_path)
            tables = []
            for page in doc:
                tables += detector.extract(page)
            doc.close()
            if len(tables) == 0:
                return None, None
            else:
                return tables
        
        # Check and read tables
        tables = extract_tables(path)        
        formatted_table_jsons = []
        if tables is not None:
            formatter = AutoTableFormatter()
            for i in tables:
                try:
                    formatted_table = formatter.extract(i)
                    df = formatted_table.df()
                    df.columns = make_unique_columns(df.columns)
                    df.columns = pd.Series(df.columns).astype(str).map(lambda x: x.strip())
                    json_structure = df.to_json(orient='records', indent=2, force_ascii=False)
                    json_structure = f'JSON Table: \njson```\n{str(json_structure)}\n```'
                    formatted_table_jsons.append(json_structure)
                except:
                    pass
        
        if len(formatted_table_jsons) == 0:
            return None
        return formatted_table_jsons
    
    def read_tables_docx(self, path):
        document = Document(path)
        tables = []
        for table in document.tables:
            df = [['' for _ in range(len(table.columns))] for _ in range(len(table.rows))]
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    if cell.text:
                        df[i][j] = cell.text
            tables.append(str(df))
        return tables
            
    def read_pdf(self, path):        
        # Read text
        with open(path, 'rb') as file:
            reader = pypdf.PdfReader(file, strict=False)
            text = ""
            for page in reader.pages:
                content = page.extract_text()
                text += content
        return text
    

    def read_txt(self, path):
        text = ""
        with open(path, 'r') as file:
            for line in file:
                text += line 
        return text
    

    def read_docx(self, path):
        text = ""
        doc = docx.api.Document(path)
        for p in doc.paragraphs:
            text += p.text + "\n"
        return text


    def read_doc(self, path):
        document = spire.doc.Document()
        document.LoadFromFile(path)
        section = document.Sections[0]
        text = ''
        for i in range(section.Paragraphs.Count):
            paragraph = section.Paragraphs[i]
            text += paragraph.Text + "\n"
        return text
    