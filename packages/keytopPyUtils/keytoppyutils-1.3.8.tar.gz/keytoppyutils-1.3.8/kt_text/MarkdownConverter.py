from io import BytesIO

import markdown
import requests
from PIL import Image
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from weasyprint import HTML

from kt_base.CommonUtils import CommonUtils
from kt_base.FileUtils import FileUtils
from kt_text.Config import Config


class MarkdownConverter:
    def __init__(self, markdown_content):
        self.markdown_content = markdown_content.replace("：\n", "：\n\n").replace("：\\*\\*\n", "：**\n\n")
        self.html_content = markdown.markdown(self.markdown_content, extensions=['tables'])
        self.html_content = self.move_p_content_to_li(self.html_content)
        FileUtils.create_paths(Config.BASE_PATH)

    def to_html(self):
        return self.html_content

    def move_p_content_to_li(html):
        soup = BeautifulSoup(html, 'html.parser')

        # 遍历所有的 <ul> 标签
        for ul in soup.find_all('ul'):
            # 遍历 <ul> 下的所有 <li> 标签
            for li in ul.find_all('li'):
                # 遍历 <li> 下的所有 <p> 标签
                for p in li.find_all('p'):
                    # 提取 <p> 标签的内容
                    p_content = "".join(p.decode_contents())
                    # 在 <li> 标签中插入 <p> 标签的内容
                    p.insert_before(BeautifulSoup(p_content, 'html.parser'))
                    # 移除 <p> 标签
                    p.decompose()
                # 移除 <li> 下的所有空白字符
                for child in li.contents:
                    if isinstance(child, NavigableString) and child.strip() == '':
                        child.extract()

        # 返回处理后的 HTML 文本
        return str(soup)
    def to_word(self):
        """
        将markdown文本转换成word，
        :return: 返回文件名
        """
        file_name = CommonUtils.generate_uuid() + ".docx"

        # 创建 Word 文档
        doc = Document()
        # 设置文档样式
        style = doc.styles['Normal']
        font = style.font
        font.name = 'SimSun'
        font.size = Pt(12)
        soup = BeautifulSoup(self.html_content, 'html.parser')
        # 处理标题和段落
        for tag in soup.find_all(['h1', 'h2', 'h3', 'h4','h5','h6', 'p', 'table','img','ul','li']):
            if tag.name == 'h1':
                # 添加一级标题
                heading = doc.add_heading(level=1)
                heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(18)
            elif tag.name == 'h2':
                # 添加二级标题
                heading = doc.add_heading(level=2)
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(16)
            elif tag.name == 'h3':
                # 添加三级标题
                heading = doc.add_heading(level=3)
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(14)
            elif tag.name == 'h4':
                # 添加四级标题
                heading = doc.add_heading(level=4)
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(12)
            elif tag.name == 'h5':
                # 添加四级标题
                heading = doc.add_heading(level=4)
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(12)
            elif tag.name == 'h6':
                # 添加四级标题
                heading = doc.add_heading(level=4)
                heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(14)
            elif tag.name == 'p':
                # 添加段落
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(tag.get_text())
                run.font.size = Pt(12)  # 设置字体大小
            elif tag.name == 'img':
                # 添加图片
                img_src = tag.get('src')
                if img_src:
                    response = requests.get(img_src)
                    if response.status_code == 200:
                        img_data = BytesIO(response.content)
                        img = Image.open(img_data)
                        original_width, original_height = img.size

                        # 设置固定的宽度
                        fixed_width = Inches(5)  # 例如，设置宽度为5英寸

                        # 计算新的高度，保持比例
                        scale_factor = fixed_width.inches / original_width
                        new_height = original_height * scale_factor

                        # 将图片数据重新加载到BytesIO对象中
                        img_data.seek(0)
                        doc.add_picture(img_data, width=fixed_width, height=Inches(new_height))
            elif tag.name == 'ul':
                # 遍历每个 <ul> 标签下的 <li> 标签
                for li in tag.find_all('li'):
                    # 创建一个新的段落
                    paragraph = doc.add_paragraph()
                    # 遍历列表项中的所有子元素
                    for child in li.children:
                        if child.name == 'strong':
                            # 添加加粗部分
                            run = paragraph.add_run(child.get_text(strip=True))
                            run.bold = True
                        else:
                            text = child.get_text(strip=True)
                            if not (text is None or text=='' or text=="\n"):
                                # 添加普通文本
                                paragraph.add_run(text)

            elif tag.name == 'table':
                # 添加表格
                table_data = []
                for row in tag.find_all('tr'):
                    cells = [cell.get_text(strip=True) for cell in row.find_all(['th', 'td'])]
                    table_data.append(cells)

                # 创建表格
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'  # 应用内置表格样式
                table.alignment = WD_TABLE_ALIGNMENT.LEFT  # 表格居中对齐

                # 设置表头样式
                for cell in table.rows[0].cells:
                    # 确保段落和运行存在
                    if not cell.paragraphs:
                        cell.add_paragraph()
                    paragraph = cell.paragraphs[0]
                    if not paragraph.runs:
                        paragraph.add_run()
                    run = paragraph.runs[0]
                    run.bold = True
                    run.font.size = Pt(12)
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    cell.vertical_alignment = 1  # 垂直居中
                    shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shading_elm)  # 设置背景色

                # 填充表格数据
                for i, row in enumerate(table_data):
                    for j, cell in enumerate(row):
                        table.cell(i, j).text = cell
                        # 确保段落和运行存在
                        if not table.cell(i, j).paragraphs:
                            table.cell(i, j).add_paragraph()
                        paragraph = table.cell(i, j).paragraphs[0]
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        table.cell(i, j).vertical_alignment = 1  # 垂直居中

                # 设置每行的高度
                for row in table.rows:
                    tr = row._tr
                    trPr = tr.get_or_add_trPr()
                    trHeight = parse_xml(r'<w:trHeight {} w:val="500" w:hRule="atLeast"/>'.format(nsdecls('w')))
                    trPr.append(trHeight)

        # 保存 Word 文档
        doc.save(Config.BASE_PATH + file_name)
        return file_name

    def to_pdf(self,style):
        """
        根据给定的样式，将markdown文本转换成PDF
        :param style: 样式内容，需要设置body、H1-H6、table等的样式，用来控制
        :return: 文件名
        """
        file_name = CommonUtils.generate_uuid()+ ".pdf";
        html_text =  style + self.html_content
        HTML(string=html_text).write_pdf(Config.BASE_PATH+ file_name)
        return file_name

str = """
# keytop数据分析报告\n###### 2024-09\n\n## 总结\n本次分析的停车场为1个，之江饭店停车场。整体来看，车场在临停管理和抵扣管理方面存在一些问题。具体表现在：\n- 异常放行金额占比为1.00%，接近警戒线，需要关注。\n- 临停异常放行数据中，出场匹配不到入场车辆占比为0.08%，未超过1%，设备运行正常。\n- 人为操作异常笔数较少，没有发现明显的设备或操作问题。\n- 月租车位占比为41.25%，未超出总车位数，但免费车位占比高达36.97%，需进一步优化。\n- 抵扣金额占整体临停应收金额较大，主要抵扣来源为商户券和西软酒店对接，部分车牌抵扣次数较多。\n\n## 一、临停部分\n### 1. 异常放行和跟车放行损失情况\n#### 1.1 整体分析\n- **临停总应收金额**：469,535元\n- **异常放行总金额**：4,710元\n- **异常放行占比**：1.00%\n- **跟车放行总金额**：0元\n- **跟车放行占比**：0.00%\n- **总损失金额**：4,710元\n- **总损失占比**：1.00%\n\n异常放行和跟车放行损失金额占比加起来为1.00%，接近警戒线，但尚未超过1%。整体来看，管理处于合格状态，但需关注异常放行的情况。\n\n### 2. 临停异常放行数据\n#### 2.1 出场匹配不到入场车辆占比\n- **车场名字**：之江饭店停车场\n- **出场车辆数**：15,403辆\n- **出场匹配不到入场车辆数**：13辆\n- **不匹配占比**：0.08%\n- **是否超过1%**：否\n\n出场匹配不到入场车辆占比为0.08%，未超过1%，设备运行正常。\n\n### 3. 人为操作异常\n#### 3.1 人为操作异常笔数\n- **总人为操作笔数**：1笔\n- **异常操作记录**：\n  - **车场**：之江饭店停车场\n  - **操作人员**：无\n  - **操作类型**：无\n  - **异常笔数**：1笔\n\n人为操作异常笔数较少，仅为1笔，未发现明显的设备或操作问题。\n\n## 二、月租数据分析\n### 1. 月租车位数超出总车位数情况\n- **车场名称**：之江饭店停车场\n- **月租车位数**：165个\n- **总车位数**：400个\n- **月租车位占比**：41.25%\n- **占比是否超过100%**：否\n\n之江饭店停车场的月租车位数未超出总车位数，占比为41.25%。\n\n### 2. 免费套餐情况\n- **车场名称**：之江饭店停车场\n- **免费车位数**：61个\n- **收费车位数**：104个\n- **免费车位占比**：36.97%\n\n之江饭店停车场的免费车位占比为36.97%，超过15%，可能存在月租管理问题。\n\n ![](http://1.94.15.153:5000/api/file/img-view/2024-11-12/png/line_chart_with_colors5.png)\n## 三、抵扣数据分析\n### 1. 抵扣金额占比情况\n- **所有车场抵扣订单总数**：8,239笔\n- **所有车场抵扣金额**：357,340元\n- **所有车场抵扣金额占比**：0%\n\n抵扣金额占整体临停应收金额较大，具体情况如下：\n\n### 2. 各个抵扣来源的抵扣金额及占比情况\n- **主要抵扣来源**：\n  - **抵扣来源**：商户券\n  - **抵扣金额**：279,310元\n  - **抵扣金额占比**：0%\n  - **抵扣来源**：西软酒店对接\n  - **抵扣金额**：64,510元\n  - **抵扣金额占比**：0%\n  - **抵扣来源**：开迈斯新能源科技\n  - **抵扣金额**：13,470元\n  - **抵扣金额占比**：0%\n\n主要抵扣来源为商户券和西软酒店对接。\n\n### 3. 减免金额较多的商户\n- **减免停车金额较多的商户**：\n  - **车场**：之江饭店停车场\n  - **商户**：朗达科技\n  - **分摊收入**：686元\n  - **减免停车金额**：980元\n  - **盈亏金额**：-294元\n  - **商户**：杭州壳研\n  - **分摊收入**：161元\n  - **减免停车金额**：160元\n  - **盈亏金额**：1元\n\n减免金额较多的商户有朗达科技和杭州壳研。\n\n### 4. 抵扣异常的车牌\n- **抵扣超过15次的车牌**：\n\n  | 车场 | 车牌号 | 抵扣次数 |\n  | --- | --- | --- |\n  | 之江饭店停车场 | 浙ACQ433 | 27 |\n  | 之江饭店停车场 | 浙ABE1207 | 20 |\n  | 之江饭店停车场 | 浙ABA3323 | 17 |\n  | 之江饭店停车场 | 浙A2T926 | 17 |\n  | 之江饭店停车场 | 浙ABP9590 | 16 |\n\n抵扣次数较多的车牌前5名如上表所示。\n\n 
"""
MarkdownConverter(str).to_word()