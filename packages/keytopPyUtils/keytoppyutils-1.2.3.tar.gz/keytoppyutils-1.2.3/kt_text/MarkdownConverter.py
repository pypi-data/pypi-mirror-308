from io import BytesIO

import markdown
import requests
from PIL import Image
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from weasyprint import HTML

from kt_base.CommonUtils import CommonUtils
from kt_base.FileUtils import FileUtils
from kt_text.Config import Config


class MarkdownConverter:
    def __init__(self, markdown_content):
        self.markdown_content = markdown_content.replace("：\n", "：\n\n").replace("：\\*\\*\n", "：**\n\n")
        self.html_content = markdown.markdown(self.markdown_content, extensions=['tables'])
        FileUtils.create_paths(Config.BASE_PATH)

    def to_html(self):
        return self.html_content

    def to_word(self):
        """
        将markdown文本转换成word，
        :return: 返回文件名
        """
        file_name = CommonUtils.generate_uuid() + ".docx"

        # 创建 Word 文档
        doc = Document()
        # 设置文档样式
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(12)
        soup = BeautifulSoup(self.html_content, 'html.parser')
        # 处理标题和段落
        for tag in soup.find_all(['h1', 'h2', 'h3', 'h4','h5','h6', 'p', 'table','img']):
            if tag.name == 'h1':
                # 添加一级标题
                heading = doc.add_heading(level=1)
                heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(18)
            elif tag.name == 'h2':
                # 添加二级标题
                heading = doc.add_heading(level=2)
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(16)
            elif tag.name == 'h3':
                # 添加三级标题
                heading = doc.add_heading(level=3)
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(14)
            elif tag.name == 'h4':
                # 添加四级标题
                heading = doc.add_heading(level=4)
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(12)
            elif tag.name == 'h5':
                # 添加四级标题
                heading = doc.add_heading(level=4)
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(12)
            elif tag.name == 'h6':
                # 添加四级标题
                heading = doc.add_heading(level=4)
                heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = heading.add_run(tag.get_text())
                run.bold = True
                run.font.size = Pt(14)
            elif tag.name == 'p':
                # 添加段落
                doc.add_paragraph(tag.get_text())
                run.font.size = Pt(12)
            elif tag.name == 'img':
                # 添加图片
                img_src = tag.get('src')
                if img_src:
                    response = requests.get(img_src)
                    if response.status_code == 200:
                        img_data = BytesIO(response.content)
                        img = Image.open(img_data)
                        original_width, original_height = img.size

                        # 设置固定的宽度
                        fixed_width = Inches(5)  # 例如，设置宽度为5英寸

                        # 计算新的高度，保持比例
                        scale_factor = fixed_width.inches / original_width
                        new_height = original_height * scale_factor

                        # 将图片数据重新加载到BytesIO对象中
                        img_data.seek(0)
                        doc.add_picture(img_data, width=fixed_width, height=Inches(new_height))
            elif tag.name == 'table':
                # 添加表格
                table_data = []
                for row in tag.find_all('tr'):
                    cells = [cell.get_text(strip=True) for cell in row.find_all(['th', 'td'])]
                    table_data.append(cells)

                # 创建表格
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'  # 应用内置表格样式
                table.alignment = WD_TABLE_ALIGNMENT.LEFT  # 表格居中对齐

                # 设置表头样式
                for cell in table.rows[0].cells:
                    # 确保段落和运行存在
                    if not cell.paragraphs:
                        cell.add_paragraph()
                    paragraph = cell.paragraphs[0]
                    if not paragraph.runs:
                        paragraph.add_run()
                    run = paragraph.runs[0]
                    run.bold = True
                    run.font.size = Pt(12)
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    cell.vertical_alignment = 1  # 垂直居中
                    shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shading_elm)  # 设置背景色

                # 填充表格数据
                for i, row in enumerate(table_data):
                    for j, cell in enumerate(row):
                        table.cell(i, j).text = cell
                        # 确保段落和运行存在
                        if not table.cell(i, j).paragraphs:
                            table.cell(i, j).add_paragraph()
                        paragraph = table.cell(i, j).paragraphs[0]
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        table.cell(i, j).vertical_alignment = 1  # 垂直居中

                # 设置每行的高度
                for row in table.rows:
                    tr = row._tr
                    trPr = tr.get_or_add_trPr()
                    trHeight = parse_xml(r'<w:trHeight {} w:val="500" w:hRule="atLeast"/>'.format(nsdecls('w')))
                    trPr.append(trHeight)

        # 保存 Word 文档
        doc.save(Config.BASE_PATH + file_name)
        return file_name

    def to_pdf(self,style):
        """
        根据给定的样式，将markdown文本转换成PDF
        :param style: 样式内容，需要设置body、H1-H6、table等的样式，用来控制
        :return: 文件名
        """
        file_name = CommonUtils.generate_uuid()+ ".pdf";
        html_text =  style + self.html_content
        HTML(string=html_text).write_pdf(Config.BASE_PATH+ file_name)
        return file_name
